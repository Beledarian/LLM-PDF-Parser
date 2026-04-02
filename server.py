from typing import Any, Sequence
import os
import unicodedata
from pypdf import PdfReader
from mcp.server.fastmcp import FastMCP
import requests
import tempfile
import pytesseract
from PIL import Image
import io
from docx import Document

def _normalize_text(text: str) -> str:
    """Normalize Unicode text to NFC form.

    Composes combining characters into single glyphs so that umlauts and
    accented letters extracted from PDFs are represented consistently.
    Example: 'a' + combining diaeresis (NFD) → 'ä' (NFC).
    """
    if not text:
        return text
    return unicodedata.normalize('NFC', text)


# Initialize the MCP server
mcp = FastMCP("PDF-Parser")

# Configure Tesseract path for Windows
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

@mcp.tool()
def read_pdf(file_path: str, pages: str = None, password: str = None, ocr: bool = True, ocr_language: str = "deu+eng") -> str:
    """
    Reads a PDF file (local path or URL) and converts it to Markdown-ish text.
    
    Args:
        file_path: The absolute path to the local PDF file OR a URL (http/https).
        pages: (Optional) A comma-separated list of page numbers to read (0-indexed). 
               Example: "0,1,5" or "0-5". If omitted, reads the entire file.
        password: (Optional) Password for decrypting the PDF if it is protected.
        ocr: (Optional) If True, also extracts text from images using Tesseract OCR.
             Useful for scanned documents. Default is True.
        ocr_language: (Optional) Tesseract language string for OCR, e.g. "deu+eng" for
                      German + English (supports umlauts and other special characters).
                      Requires the corresponding Tesseract language packs to be installed.
                      Default is "deu+eng".
    
    Note: This implementation uses pypdf (pure Python) instead of pymupdf due to 
          compilation issues on Windows ARM64. The output is plain text extracted 
          from the PDF, not as sophisticated as pymupdf4llm's Markdown output.
    """
    temp_pdf_path = None
    reader = None
    
    try:
        # Handle URL inputs
        if file_path.startswith(('http://', 'https://')):
            try:
                response = requests.get(file_path, stream=True)
                response.raise_for_status()
                
                # Create a temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    for chunk in response.iter_content(chunk_size=8192):
                        tmp.write(chunk)
                    temp_pdf_path = tmp.name
                
                # Use the temp file path for processing
                target_path = temp_pdf_path
            except Exception as e:
                return f"Error downloading PDF from URL: {str(e)}"
        else:
            # Handle local file inputs
            if not os.path.exists(file_path):
                return f"Error: File not found at {file_path}"
            target_path = file_path

        # Open the PDF with pypdf
        try:
            reader = PdfReader(target_path)
        except Exception as e:
            return f"Error opening PDF: {str(e)}"

        # Handle password-protected PDFs
        if reader.is_encrypted:
            if password:
                try:
                    reader.decrypt(password)
                except Exception as e:
                    return f"Error: Failed to decrypt PDF with provided password: {str(e)}"
            else:
                return "Error: PDF is password protected but no password was provided."

        # Parse page selection if provided
        page_indices = None
        if pages:
            page_indices = []
            for part in pages.split(','):
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    page_indices.extend(range(start, end + 1))
                else:
                    page_indices.append(int(part))
        
        # Extract text from pages
        extracted_text = []
        total_pages = len(reader.pages)
        
        pages_to_process = page_indices if page_indices else range(total_pages)
        
        for idx in pages_to_process:
            if page_indices and (idx < 0 or idx >= total_pages):
                extracted_text.append(f"# Page {idx + 1}\n\nError: Page index out of range")
                continue
                
            page = reader.pages[idx]
            page_content = []
            
            # Extract embedded text
            text = page.extract_text()
            if text and text.strip():
                page_content.append(_normalize_text(text))
            
            # OCR images if requested
            if ocr:
                try:
                    images = page.images
                    for img_idx, image in enumerate(images):
                        try:
                            # Get image data
                            img_data = image.data
                            img = Image.open(io.BytesIO(img_data))
                            
                            # Run OCR
                            ocr_text = pytesseract.image_to_string(img, lang=ocr_language)
                            if ocr_text and ocr_text.strip():
                                page_content.append(f"\n[OCR from image {img_idx + 1}]\n{_normalize_text(ocr_text).strip()}")
                        except Exception as img_e:
                            page_content.append(f"\n[OCR Error for image {img_idx + 1}: {str(img_e)}]")
                except Exception as ocr_e:
                    page_content.append(f"\n[OCR extraction error: {str(ocr_e)}]")
            
            if page_content:
                extracted_text.append(f"# Page {idx + 1}\n\n" + "\n".join(page_content))
            else:
                extracted_text.append(f"# Page {idx + 1}\n\n[No text content found]")
        
        return "\n\n---\n\n".join(extracted_text)

    except Exception as e:
        return f"Error parsing PDF: {str(e)}"
    
    finally:
        # Clean up temporary file if it was created
        if temp_pdf_path and os.path.exists(temp_pdf_path):
            try:
                os.remove(temp_pdf_path)
            except OSError:
                pass

@mcp.tool()
def read_docx(file_path: str, ocr: bool = True, ocr_language: str = "deu+eng") -> str:
    """
    Reads a DOCX file (local path or URL) and converts it to Markdown-ish text.
    
    Args:
        file_path: The absolute path to the local DOCX file OR a URL (http/https).
        ocr: (Optional) If True, also extracts text from embedded images using Tesseract OCR.
             Default is True.
        ocr_language: (Optional) Tesseract language string for OCR, e.g. "deu+eng" for
                      German + English (supports umlauts and other special characters).
                      Default is "deu+eng".
    
    Returns:
        Markdown-formatted text extracted from the DOCX file.
    """
    temp_docx_path = None
    
    try:
        # Handle URL inputs
        if file_path.startswith(('http://', 'https://')):
            try:
                response = requests.get(file_path, stream=True)
                response.raise_for_status()
                
                # Create a temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    for chunk in response.iter_content(chunk_size=8192):
                        tmp.write(chunk)
                    temp_docx_path = tmp.name
                
                target_path = temp_docx_path
            except Exception as e:
                return f"Error downloading DOCX from URL: {str(e)}"
        else:
            # Handle local file inputs
            if not os.path.exists(file_path):
                return f"Error: File not found at {file_path}"
            target_path = file_path

        # Open the DOCX file
        try:
            doc = Document(target_path)
        except Exception as e:
            return f"Error opening DOCX: {str(e)}"

        extracted_content = []
        
        # Extract paragraphs with style detection
        for para in doc.paragraphs:
            text = _normalize_text(para.text.strip())
            if not text:
                continue
            
            # Detect heading styles and format accordingly
            style_name = para.style.name.lower() if para.style else ""
            if "heading 1" in style_name:
                extracted_content.append(f"# {text}")
            elif "heading 2" in style_name:
                extracted_content.append(f"## {text}")
            elif "heading 3" in style_name:
                extracted_content.append(f"### {text}")
            elif "heading 4" in style_name:
                extracted_content.append(f"#### {text}")
            elif "title" in style_name:
                extracted_content.append(f"# {text}")
            elif "list" in style_name:
                extracted_content.append(f"- {text}")
            else:
                extracted_content.append(text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_content = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                table_content.append("| " + " | ".join(cells) + " |")
                
                # Add header separator after first row
                if row_idx == 0:
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    table_content.append(separator)
            
            if table_content:
                extracted_content.append(f"\n**Table {table_idx + 1}:**\n" + "\n".join(table_content))
        
        # OCR images if requested
        if ocr:
            try:
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                
                # Access embedded images through the document's relationships
                for rel_id, rel in doc.part.rels.items():
                    if "image" in rel.reltype:
                        try:
                            image_part = rel.target_part
                            img_data = image_part.blob
                            img = Image.open(io.BytesIO(img_data))
                            
                            ocr_text = pytesseract.image_to_string(img, lang=ocr_language)
                            if ocr_text and ocr_text.strip():
                                extracted_content.append(f"\n[OCR from embedded image]\n{_normalize_text(ocr_text).strip()}")
                        except Exception as img_e:
                            extracted_content.append(f"\n[OCR Error for image: {str(img_e)}]")
            except Exception as ocr_e:
                extracted_content.append(f"\n[OCR extraction error: {str(ocr_e)}]")
        
        if extracted_content:
            return "\n\n".join(extracted_content)
        else:
            return "[No text content found in DOCX]"
    
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"
    
    finally:
        # Clean up temporary file if it was created
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except OSError:
                pass

if __name__ == "__main__":
    mcp.run()
